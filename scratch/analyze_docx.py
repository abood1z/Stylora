
import docx

def analyze_full_doc(file_path):
    doc = docx.Document(file_path)
    print("--- FIRST 100 LINES ---")
    for i, para in enumerate(doc.paragraphs[:100]):
        if para.text.strip():
            print(f"[{i}] {para.text}")
    
    print("\n--- TABLES ---")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for row in table.rows:
            print(" | ".join(cell.text.strip() for cell in row.cells))

if __name__ == "__main__":
    analyze_full_doc(r"d:\project_folder\documentationstylora.docx")
