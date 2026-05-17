
import docx
import sys

# Set stdout to utf-8 just in case
sys.stdout.reconfigure(encoding='utf-8')

def analyze_more(file_path):
    doc = docx.Document(file_path)
    print("--- PARAGRAPHS 0-30 ---")
    for i, para in enumerate(doc.paragraphs[:30]):
        print(f"[{i}] {para.text}")
    
    print("\n--- TABLE 0 ---")
    if doc.tables:
        table = doc.tables[0]
        for row in table.rows:
            print(" | ".join(cell.text.strip() for cell in row.cells))

if __name__ == "__main__":
    analyze_more(r"d:\project_folder\documentationstylora.docx")
