
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

def append_chapters_3_4(file_path):
    doc = docx.Document(file_path)
    
    # Chapter 3
    doc.add_page_break()
    p = doc.add_paragraph("CHAPTER 3")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.bold = True
    
    p = doc.add_paragraph("METHODOLOGY")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.bold = True
    
    doc.add_heading("3.1 Research Methodology", level=2)
    doc.add_paragraph("This project follows a systematic research approach combining qualitative requirements analysis and quantitative AI model evaluation. The study begins with an analysis of user pain points in wardrobe management, followed by the selection of appropriate computer vision techniques to address these challenges.")
    
    doc.add_heading("3.2 Software Development Lifecycle (Waterfall Model)", level=2)
    doc.add_paragraph("For this graduation project, the 'Waterfall Model' was selected due to its structured and sequential nature, which ensures that requirements are fully defined before development begins. The stages include:")
    doc.add_paragraph("1. Requirements Analysis: Defining the core features of Stylora.", style='List Bullet')
    doc.add_paragraph("2. System Design: Creating UML diagrams and database schemas.", style='List Bullet')
    doc.add_paragraph("3. Implementation: Coding the mobile app and training the CNN model.", style='List Bullet')
    doc.add_paragraph("4. Testing: Validating the accuracy of AI matching and UI responsiveness.", style='List Bullet')
    
    doc.add_heading("3.3 Data Acquisition", level=2)
    doc.add_paragraph("The AI model's training data was sourced from public fashion datasets (e.g., DeepFashion) and locally collected garment images to ensure the model can recognize locally preferred styles and cloth types in Jordan.")
    
    # Chapter 4
    doc.add_page_break()
    p = doc.add_paragraph("CHAPTER 4")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.bold = True
    
    p = doc.add_paragraph("ANALYSIS AND DESIGN")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.bold = True
    
    doc.add_heading("4.1 Requirements Determination", level=2)
    doc.add_paragraph("The system's requirements have been categorized into functional and non-functional requirements to ensure a comprehensive development scope.")
    
    doc.add_heading("4.1.1 Functional Requirements", level=3)
    doc.add_paragraph("FR1: The system shall allow users to upload images of their garments via the mobile camera.", style='List Bullet')
    doc.add_paragraph("FR2: The system shall utilize a CNN to classify the garment type and color automatically.", style='List Bullet')
    doc.add_paragraph("FR3: The system shall provide automated outfit recommendations based on the user's digital closet.", style='List Bullet')
    doc.add_paragraph("FR4: The system shall provide a dashboard for merchants to manage their garment store.", style='List Bullet')
    
    doc.add_heading("4.2 Logical Design", level=2)
    
    doc.add_heading("4.2.1 Use Case Diagram", level=3)
    doc.add_paragraph("The following diagram illustrates the interaction between the 'Customer' and 'Merchant' actors and the Stylora system.")
    doc.add_paragraph("[IMAGE PLACEHOLDER: Use Case Diagram]")
    
    doc.add_heading("4.2.2 Sequence Diagram", level=3)
    doc.add_paragraph("This diagram details the sequence of operations for the AI recommendation flow.")
    doc.add_paragraph("[IMAGE PLACEHOLDER: Sequence Diagram]")
    
    doc.add_heading("4.2.3 Flow Chart", level=3)
    doc.add_paragraph("The logic flow of the AI matching algorithm is detailed below.")
    doc.add_paragraph("[IMAGE PLACEHOLDER: Flow Chart]")
    
    doc.add_heading("4.2.4 Entity Relationship Diagram (ERD)", level=3)
    doc.add_paragraph("The database schema for managing users, closet items, and products is represented in the following ERD.")
    doc.add_paragraph("[IMAGE PLACEHOLDER: ER Diagram]")
    
    doc.add_heading("4.3 User Interface (Screenshots)", level=2)
    doc.add_paragraph("The initial prototypes of the Stylora application interfaces are shown below:")
    doc.add_paragraph("1. Splash Screen & Login: Smooth entry with Stylora branding.", style='List Bullet')
    doc.add_paragraph("2. Virtual Closet: Visual grid display of user garments.", style='List Bullet')
    doc.add_paragraph("3. AI Suggestion View: Matching results with 'Wear' or 'Shop' options.", style='List Bullet')
    doc.add_paragraph("4. Trader Dashboard: Inventory and sales tracking for merchants.", style='List Bullet')
    
    doc.save(file_path)
    print("Appended Chapter 3 and 4 successfully.")

if __name__ == "__main__":
    append_chapters_3_4(r"d:\project_folder\documentationstylora_corrected.docx")
