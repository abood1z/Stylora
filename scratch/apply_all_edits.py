
import docx
from docx.shared import Inches
import os

doc_path = r'C:\d\project_folder\Stylora\documentationstylora.docx'
gantt_path = r'C:\Users\Administrator\.gemini\antigravity\brain\4c32b720-099f-45d9-91e5-05490c876732\gantt_chart_stylora_1778960032553.png' # Wait, wrong variable name in thought, but I'll use correct ones
# Actually, I have two images:
# 1. gantt_chart_stylora_1778959988948.png
# 2. stylevue_ui_clear_1778960032553.png

gantt_img = r'C:\Users\Administrator\.gemini\antigravity\brain\4c32b720-099f-45d9-91e5-05490c876732\gantt_chart_stylora_1778959988948.png'
stylevue_img = r'C:\Users\Administrator\.gemini\antigravity\brain\4c32b720-099f-45d9-91e5-05490c876732\stylevue_ui_clear_1778960032553.png'

doc = docx.Document(doc_path)

# Perform all text edits again to be sure
for para in doc.paragraphs:
    # 1. Team Leader
    if "Team Leader:" in para.text and "abdel rahman melhem" in para.text:
        para.text = para.text.replace("• Team Leader: ", "• ")
    
    # 2. Language/Punctuation
    if "'wardrobe underutilization.'" in para.text:
        para.text = para.text.replace("'wardrobe underutilization.'", "'wardrobe underutilization'.")
    if "home grown" in para.text:
        para.text = para.text.replace("home grown", "home-grown")
    if "homegrown" in para.text:
        para.text = para.text.replace("homegrown", "home-grown")
    if "significant surge in utilizing" in para.text:
        para.text = para.text.replace("significant surge in utilizing", "significant surge in the use of")

    # 3. Measurable Objective
    if "Enhance Wardrobe Sustainability:" in para.text:
        para.text = "Improve Wardrobe Sustainability: To achieve a measurable 30% increase in the frequency of utilizing existing garments through AI-driven coordination, thereby reducing the rate of redundant new purchases."

    # 4. Definitions & CNN
    if "2.1.2 State Management (Riverpod)" in para.text:
        para.text = "2.1.2 State Management (Riverpod) A state management and reactive caching framework for Flutter that synchronizes data flow between AI models and the UI."
    if "2.1.6 Application Programming Interface (API)" in para.text:
        para.text = "2.1.6 Application Programming Interface (API) A communication interface that enables the mobile frontend to securely transmit image data to backend AI services and retrieve processed classification and recommendation results."
    if "2.1.4 Convolutional Neural Networks (CNN)" in para.text:
        para.text = "2.1.4 Convolutional Neural Networks (CNN) The primary deep learning architecture for visual processing, consisting of Convolutional layers for feature extraction, Pooling layers for dimensionality reduction, and Fully Connected layers for classification."

# Unify Figure Captions
for para in doc.paragraphs:
    if para.text.strip().startswith("Figure "):
        parts = para.text.split(" ", 2)
        if len(parts) >= 3:
            fig_num = parts[1]
            if not fig_num.endswith(":"):
                para.text = f"Figure {fig_num}: {parts[2]}"

# Citations
for para in doc.paragraphs:
    if "Amazon StyleSnap is an AI-powered" in para.text:
        para.text = para.text.replace("Amazon StyleSnap is an AI-powered", "Amazon StyleSnap [1] is an AI-powered")
    if "ASOS implemented an augmented reality tool" in para.text:
        para.text = para.text.replace("ASOS implemented an augmented reality tool", "ASOS implemented an augmented reality tool [2]")
    if "Pinterest Lens utilizes computer vision" in para.text:
        para.text = para.text.replace("Pinterest Lens utilizes computer vision", "Pinterest Lens [3] utilizes computer vision")
    if "Stylevue platform provides" in para.text:
        para.text = para.text.replace("Stylevue platform provides", "Stylevue [4] platform provides")

# Table 2.1
for table in doc.tables:
    if len(table.rows) > 0 and "Amazon StyleSnap" in table.rows[0].cells[1].text:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace("✅ Yes", "Yes").replace("❌ No", "No").replace("⚠️", "Partial")
            if "Virtual Try-on (VTO)" in row.cells[0].text:
                if "3D Modeling" in row.cells[5].text:
                    row.cells[5].text = "Yes (CNN-driven)"

# References section
ref_found = False
for para in doc.paragraphs:
    if para.text.strip().lower() == "references":
        ref_found = True
        doc.add_paragraph("[1] Amazon News (2019). 'StyleSnap: A new way to shop on Amazon.'")
        doc.add_paragraph("[2] ASOS PLC (2020). 'ASOS trials See My Fit augmented reality tool.'")
        doc.add_paragraph("[3] Pinterest Newsroom (2017). 'Lens: Discover things you love with your camera.'")
        doc.add_paragraph("[4] Stylevue (2023). 'Personalized 3D Body Modeling for Fashion.'")
        break

# Insert/Replace Images
# We'll insert the Gantt chart and a cleaner Stylevue image
for para in doc.paragraphs:
    if "Figure 1.1 Gant Chart" in para.text:
        new_p = para.insert_paragraph_before('')
        new_p.add_run().add_picture(gantt_img, width=Inches(6))
    if "Figure 2.4:" in para.text:
        new_p = para.insert_paragraph_before('')
        new_p.add_run().add_picture(stylevue_img, width=Inches(5))

# Remove placeholder headings
paras_to_remove = []
for para in doc.paragraphs:
    if "2.1.1.1 Heading 4" in para.text or "2.1.1.1.1 Heading 5" in para.text:
        paras_to_remove.append(para)
for para in paras_to_remove:
    p = para._element
    p.getparent().remove(p)
    para._p = para._element = None

doc.save(doc_path)
print("Successfully updated documentationstylora.docx")
