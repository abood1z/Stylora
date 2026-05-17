
import docx
from docx.shared import Pt

def apply_corrections(input_path, output_path):
    doc = docx.Document(input_path)
    
    # 1. Remove "Team Leader" label
    # Para 12 is in the metadata/beginning
    for para in doc.paragraphs[:30]:
        if "Team Leader:" in para.text:
            para.text = para.text.replace("Team Leader: ", "")
            print("Fixed Team Leader label.")

    # 2. Grammar/Typos
    # - Chapter 1.1: point outside quotes
    # - Chapter 1.4: home-grown
    # - Chapter 2: surge in utilizing -> use of
    # - Sustainability goal
    # - CNN description
    
    for para in doc.paragraphs:
        t = para.text
        
        # Chapter 1.1
        if "wardrobe underutilization.'" in t:
            para.text = t.replace("wardrobe underutilization.'", "wardrobe underutilization'.")
            print("Fixed Chapter 1.1 punctuation.")
            
        # Chapter 2 intro
        if "significant surge in utilizing" in t:
            para.text = t.replace("significant surge in utilizing", "significant surge in the use of")
            print("Fixed Chapter 2 introduction.")
            
        # Sustainability goal (Measurable)
        if "Enhance Wardrobe Sustainability: To help users" in t:
            para.text = "Enhance Wardrobe Sustainability: Promote wardrobe sustainability by increasing the utilization rate of existing items by 25% and reducing redundant fashion purchases through intelligent digital tracking."
            print("Fixed Sustainability objective.")

        # CNN (Adding layers)
        if "detecting style features from images captured via the mobile camera." in t and "Convolutional Neural Networks" in t:
            para.text = "2.1.4 Convolutional Neural Network (CNN): A class of deep neural networks commonly applied to analyzing visual imagery. In Stylora, CNNs are responsible for classifying garments and consist of Convolutional Layers, Pooling Layers, and Fully Connected Layers for feature extraction and final classification."
            print("Fixed CNN definition.")

        # Riverpod / API balancing
        if "2.1.2 State Management (Riverpod)" in t:
            para.text = "2.1.2 State Management (Riverpod): A reactive state-management framework for Flutter that ensures compile-time safety and modularity, allowing Stylora to manage data flow efficiently across features."
            print("Fixed Riverpod definition.")
        if "2.1.6 Application Programming Interface (API)" in t:
            para.text = "2.1.6 Application Programming Interface (API): A set of protocols that allow Stylora to communicate with external services, enabling image data transmission and result retrieval from AI models."
            print("Fixed API definition.")

        # Figure 1.1 numbering (adding colon)
        if "Figure 1.1 Gant Chart" in t:
            para.text = t.replace("Figure 1.1 Gant Chart", "Figure 1.1: Gantt Chart")
            print("Fixed Figure 1.1 label.")
            
        # Remove placeholders
        if "heading 4" in t.lower() or "heading 5" in t.lower() or "first paragraph following a heading" in t.lower():
            # Only if it's not the TOC (TOC has tab leaders like \t)
            if "\t" not in t:
                para.text = "" # Clear placeholder
                print(f"Cleared placeholder: {t[:20]}...")

    # 3. Table 2.1 (Table 0)
    if doc.tables:
        table = doc.tables[0]
        # Skip header?
        for row_idx, row in enumerate(table.rows):
            if row_idx == 0: continue # Header
            for cell_idx, cell in enumerate(row.cells):
                # Replace icons
                cell.text = cell.text.replace("✅ Yes", "Yes")
                cell.text = cell.text.replace("❌ No", "No")
                cell.text = cell.text.replace("⚠️ Limited", "Partial")
                cell.text = cell.text.replace("⚠️ Recommendations", "Yes (Recs)")
                cell.text = cell.text.replace("⚠️ Suggested Fit", "Partial")
                
                # Fix Stylora Virtual Try-on (contradiction)
                if "Stylora" in table.rows[0].cells[cell_idx].text:
                    if row_idx == 2: # Cell for VTO
                        cell.text = "No (Focus on Classification and Matching)"
        print("Updated Table 2.1 indices and symbols.")

    # 4. References (Add citations paragraph)
    # Search for Chapter 2.2 end
    found_2_2 = False
    for i, para in enumerate(doc.paragraphs):
        if "2.2.4 Advanced Body Modeling" in para.text:
            found_2_2 = True
        if found_2_2 and "2.3 Comparison" in para.text:
            # Insert before this paragraph
            new_para = para.insert_paragraph_before("Information regarding global platforms was retrieved from official documentations (Amazon StyleSnap, 2019; ASOS See My Fit, 2020; Pinterest, 2017).")
            print("Added citations to Chapter 2.2.")
            break

    doc.save(output_path)
    print(f"Saved corrected document to {output_path}")

if __name__ == "__main__":
    apply_corrections(r"d:\project_folder\documentationstylora.docx", r"d:\project_folder\documentationstylora_corrected.docx")
