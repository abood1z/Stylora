
import docx
from docx.shared import Inches
import os

doc_path = r'C:\d\project_folder\Stylora\documentationstylora.docx'
gantt_path = r'C:\Users\Administrator\.gemini\antigravity\brain\4c32b720-099f-45d9-91e5-05490c876732\gantt_chart_stylora_1778959988948.png'

doc = docx.Document(doc_path)

# 1. Remove Team Leader mention
for para in doc.paragraphs:
    if "Team Leader:" in para.text and "abdel rahman melhem" in para.text:
        para.text = para.text.replace("• Team Leader: ", "• ")

# 2. Language and spelling
for para in doc.paragraphs:
    # Chapter 1.1 punctuation
    if "'wardrobe underutilization.'" in para.text:
        para.text = para.text.replace("'wardrobe underutilization.'", "'wardrobe underutilization'.")
    
    # Chapter 1.4 Unify home-grown (ensure hyphen)
    if "home grown" in para.text:
        para.text = para.text.replace("home grown", "home-grown")
    if "homegrown" in para.text:
        para.text = para.text.replace("homegrown", "home-grown")

    # Chapter 2 surge
    if "significant surge in utilizing" in para.text:
        para.text = para.text.replace("significant surge in utilizing", "significant surge in the use of")

# 3. Definitions balance and CNN layers
for para in doc.paragraphs:
    # Riverpod
    if "2.1.2 State Management (Riverpod)" in para.text:
        para.text = "2.1.2 State Management (Riverpod) A state management and reactive caching framework for Flutter that synchronizes data flow between AI models and the UI."
    # API
    if "2.1.6 Application Programming Interface (API)" in para.text:
        para.text = "2.1.6 Application Programming Interface (API) A communication interface that enables the mobile frontend to securely transmit image data to backend AI services and retrieve processed classification and recommendation results."
    # CNN
    if "2.1.4 Convolutional Neural Networks (CNN)" in para.text:
        para.text = "2.1.4 Convolutional Neural Networks (CNN) The primary deep learning architecture for visual processing, consisting of Convolutional layers for feature extraction, Pooling layers for dimensionality reduction, and Fully Connected layers for classification."

# 4. Remove placeholder headings (2.1.1.1 and 2.1.1.1.1)
# Note: I need to be careful with TOC. I'll just remove paragraphs that match the pattern exactly.
paras_to_remove = []
for para in doc.paragraphs:
    if "2.1.1.1 Heading 4" in para.text or "2.1.1.1.1 Heading 5" in para.text:
        paras_to_remove.append(para)

# Remove them from the document structure (workaround for python-docx)
for para in paras_to_remove:
    p = para._element
    p.getparent().remove(p)
    para._p = para._element = None

# 5. Measurable Objective (Chapter 1.3)
for para in doc.paragraphs:
    if "Enhance Wardrobe Sustainability:" in para.text:
        para.text = para.text.replace("Enhance Wardrobe Sustainability: To help users maximize the utility of their current clothes by providing smart styling suggestions, reducing unnecessary purchases.",
                                       "Improve Wardrobe Sustainability: To achieve a measurable 30% increase in the frequency of utilizing existing garments through AI-driven coordination, thereby reducing the rate of redundant new purchases.")

# 6. Figure Captions unification
for para in doc.paragraphs:
    if para.text.startswith("Figure "):
        # Ensure "Figure X.X: "
        if ":" not in para.text.split(" ")[2]: # Check if the number part doesn't have a colon
             parts = para.text.split(" ", 2)
             if len(parts) >= 3:
                 fig_num = parts[1]
                 if not fig_num.endswith(":"):
                     para.text = f"Figure {fig_num}: {parts[2]}"

# 7. Citations and References
# Adding [1], [2], etc. to Chapter 2.2
for para in doc.paragraphs:
    if "Amazon StyleSnap is an AI-powered" in para.text:
        para.text = para.text.replace("Amazon StyleSnap is an AI-powered", "Amazon StyleSnap [1] is an AI-powered")
    if "ASOS implemented an augmented reality tool" in para.text:
        para.text = para.text.replace("ASOS implemented an augmented reality tool", "ASOS implemented an augmented reality tool [2]")
    if "Pinterest Lens utilizes computer vision" in para.text:
        para.text = para.text.replace("Pinterest Lens utilizes computer vision", "Pinterest Lens [3] utilizes computer vision")
    if "Stylevue platform provides" in para.text:
        para.text = para.text.replace("Stylevue platform provides", "Stylevue [4] platform provides")

# Add a References section if it exists, or append to it
ref_header_found = False
for para in doc.paragraphs:
    if para.text.strip().lower() == "references":
        ref_header_found = True
        # Add sources after this
        doc.add_paragraph("[1] Amazon News (2019). 'StyleSnap: A new way to shop on Amazon.'")
        doc.add_paragraph("[2] ASOS PLC (2020). 'ASOS trials See My Fit augmented reality tool.'")
        doc.add_paragraph("[3] Pinterest Newsroom (2017). 'Lens: Discover things you love with your camera.'")
        doc.add_paragraph("[4] Stylevue (2023). 'Personalized 3D Body Modeling for Fashion.'")
        break

# 8. Table 2.1 edits
for table in doc.tables:
    # Identify if it's Table 2.1 (based on headers)
    if len(table.rows) > 0 and "Amazon StyleSnap" in table.rows[0].cells[1].text:
        for row in table.rows:
            for cell in row.cells:
                # Replace icons with text
                if "✅ Yes" in cell.text:
                    cell.text = cell.text.replace("✅ Yes", "Yes")
                if "❌ No" in cell.text:
                    cell.text = cell.text.replace("❌ No", "No")
                if "⚠️" in cell.text:
                    cell.text = cell.text.replace("⚠️", "Partial")
                
                # Fix Stylora VTO
                if "Stylora (Our Project)" in table.rows[0].cells[5].text:
                     # Row for VTO
                     if "Virtual Try-on (VTO)" in row.cells[0].text:
                         if "3D Modeling" in row.cells[5].text:
                             row.cells[5].text = "Yes (CNN-driven)"

# 9. Insert Gantt Chart Image
for para in doc.paragraphs:
    if "Figure 1.1 Gant Chart" in para.text:
        # Insert image BEFORE this caption or in a new paragraph before it
        new_p = para.insert_paragraph_before('')
        run = new_p.add_run()
        if os.path.exists(gantt_path):
            run.add_picture(gantt_path, width=Inches(6))
        else:
            new_p.text = "[Gantt Chart Image Placeholder]"

# Save the document
output_path = r'C:\d\project_folder\Stylora\documentationstylora_edited.docx'
doc.save(output_path)
print(f"Saved edited document to {output_path}")
