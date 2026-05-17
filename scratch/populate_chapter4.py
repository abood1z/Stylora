
import docx
from docx.shared import Inches
import os

doc_path = r'C:\d\project_folder\Stylora\documentationstylora.docx'
img_dir = r'C:\Users\Administrator\.gemini\antigravity\brain\4c32b720-099f-45d9-91e5-05490c876732'

# Images
usecase_img = os.path.join(img_dir, 'stylora_use_case_diagram_1778960132371.png')
sequence_img = os.path.join(img_dir, 'stylora_sequence_diagram_1778960147497.png')
dfd_img = os.path.join(img_dir, 'stylora_dfd_diagram_1778960167477.png')
erd_img = os.path.join(img_dir, 'stylora_erd_diagram_1778960183132.png')

doc = docx.Document(doc_path)

# Dictionary of content
content = {
    "4.1 Analysis": "The analysis phase focuses on defining the user needs and system specifications to ensure the Stylora platform addresses real-world challenges in the fashion industry.",
    "4.1.1 Requirements Determination": "The requirements for Stylora were gathered using a hybrid approach. Semi-structured interviews were conducted with local fashion retailers in Jordan to understand digital transformation barriers. Additionally, a quantitative survey with over 100 consumers identified persistent challenges in wardrobe coordination and online shopping uncertainty. These findings formed the basis for the project's functional priorities.",
    "4.1.2 System's Requirements": (
        "Functional Requirements (FR):\n"
        "- FR1: User Registration & Authentication: Secure access via Firebase Auth.\n"
        "- FR2: AI Vision Pipeline: Automated classification of garments (Category/Color) using CNN and YOLO.\n"
        "- FR3: Virtual Closet Management: Storage and organization of digitized garments.\n"
        "- FR4: Outfit Recommendation Engine: Algorithmically generated suggestions based on color theory.\n"
        "- FR5: Virtual Try-on (VTO): Visualization of garments on a digital profile.\n"
        "- FR6: Trader Marketplace: Tools for local vendors to list and manage inventory.\n\n"
        "Non-Functional Requirements (NFR):\n"
        "- NFR1 (Performance): AI processing and UI updates should occur within 2-3 seconds.\n"
        "- NFR2 (Usability): A clean, intuitive Flutter-based UI for high user engagement.\n"
        "- NFR3 (Scalability): Ability to handle increasing user data via cloud-native Firebase architecture."
    ),
    "4.2 Design": "The design phase translates analysis requirements into technical blueprints, covering logical interactions, data structures, and the physical architecture of the application.",
    "4.2.1 Logical Design": "The logical design describes the internal reasoning and data flow of Stylora without focusing on specific hardware constraints.",
    "4.2.1.1 Use Case Diagram": "The Use Case Diagram illustrates the interactions between the primary actors (Customer and Trader) and the system's core functionalities, highlighting the AI processing and marketplace activities.",
    "4.2.1.2 Sequence Diagrams": "The Sequence Diagram details the chronological flow of events during the garment analysis process, showing how data travels from the user's camera to the AI backend and into the cloud database.",
    "4.2.1.3 Data Flow Diagram": "The DFD represents how fashion data (images, attributes, and user preferences) is processed and stored across the system's various modules.",
    "4.2.1.4 Entity Relationship Diagram": "The ERD defines the data entities (User, ClosetItem, Outfit, Product) and their relationships within the document-oriented Firestore database.",
    "4.2.2 Physical Design": "The physical design follows a cloud-native, serverless architecture. The frontend is built with Flutter (using Riverpod for state management), while the backend is powered by Google Firebase. AI models are deployed as specialized microservices, ensuring modularity and performance. The database layer is physically implemented using Firestore collections, optimized for real-time mobile access."
}

# Helper to find paragraph by text
def find_para(text):
    for i, p in enumerate(doc.paragraphs):
        if text.strip() in p.text.strip():
            return i, p
    return -1, None

# Append content systematically
current_index = 0
for title, text in content.items():
    idx, p = find_para(title)
    if p:
        # Paragraph exists as a header, add content after it
        # Clear existing text if it was just the header number or placeholder
        if p.text.strip() == title:
            # Add description
            new_p = doc.add_paragraph(text)
            # Move it to after the header
            # (In python-docx we append or insert. Since they already exist in order, we can find and update)
            pass 
        else:
            # It's a header with text, just append the detailed text after it
            new_p = doc.add_paragraph(text)
    else:
        # Header doesn't exist, create it (shouldn't happen based on user prompt)
        doc.add_heading(title, level=2)
        doc.add_paragraph(text)

# Specialized Insertion with Images
def add_section_with_image(title, text, img_path):
    idx, p = find_para(title)
    if p:
        # Clear paragraph text after the title or append
        # For simplicity in this script, we'll append to the end of the document if we can't find clear insertion points,
        # but the user structure is at the end of the file based on the previous dump.
        pass

# Let's use a more direct approach: Search and Replace placeholders if they exist, or just append to end of specific sections.
# Looking at the dump, CHAPTER 4 starts at P241 (in the text dump logic).

# We will recreate Chapter 4 content
# First, let's find where CHAPTER 4 starts
start_idx = -1
for i, p in enumerate(doc.paragraphs):
    if "CHAPTER 4 ANALYSIS AND DESIGN" in p.text:
        start_idx = i
        break

if start_idx != -1:
    # Clear subsequent paragraphs that might be placeholders until CHAPTER 5
    end_idx = -1
    for i in range(start_idx + 1, len(doc.paragraphs)):
        if "CHAPTER 5" in doc.paragraphs[i].text:
            end_idx = i
            break
    
    # If we found CHAPTER 4, we will rebuild it
    # For now, let's just find the specific subheadings and add content
    for title, text in content.items():
        for p in doc.paragraphs[start_idx:]:
            if title in p.text and len(p.text) < len(title) + 5: # It's likely the header
                # Add text after this header
                new_p = p.insert_paragraph_before('') # Placeholder to move
                p.add_run("\n" + text)
                
                # Add images for specific sections
                if "4.2.1.1" in title and os.path.exists(usecase_img):
                    p.add_run().add_break()
                    p.add_run().add_picture(usecase_img, width=Inches(5))
                    p.add_run("\nFigure 4.1: Stylora Use Case Diagram")
                if "4.2.1.2" in title and os.path.exists(sequence_img):
                    p.add_run().add_break()
                    p.add_run().add_picture(sequence_img, width=Inches(5))
                    p.add_run("\nFigure 4.2: Stylora Sequence Diagram")
                if "4.2.1.3" in title and os.path.exists(dfd_img):
                    p.add_run().add_break()
                    p.add_run().add_picture(dfd_img, width=Inches(5))
                    p.add_run("\nFigure 4.3: Stylora Data Flow Diagram")
                if "4.2.1.4" in title and os.path.exists(erd_img):
                    p.add_run().add_break()
                    p.add_run().add_picture(erd_img, width=Inches(5))
                    p.add_run("\nFigure 4.4: Stylora Entity Relationship Diagram")
                break

doc.save(doc_path)
print("Updated Chapter 4 with detailed Analysis and Design content.")
