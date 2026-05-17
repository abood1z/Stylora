
import docx

doc = docx.Document(r'C:\d\project_folder\Stylora\documentationstylora.docx')

print("--- Paragraphs ---")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"P{i}: {para.text}")

print("\n--- Tables ---")
for i, table in enumerate(doc.tables):
    print(f"\nTable {i}:")
    for r, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"R{r}: {' | '.join(cells)}")
