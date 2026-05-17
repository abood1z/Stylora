
import docx
import sys

# Set output to utf-8
sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document(r'C:\d\project_folder\Stylora\documentationstylora.docx')

with open(r'C:\d\project_folder\Stylora\scratch\docx_content.txt', 'w', encoding='utf-8') as f:
    f.write("--- Paragraphs ---\n")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            f.write(f"P{i}: {para.text}\n")

    f.write("\n--- Tables ---\n")
    for i, table in enumerate(doc.tables):
        f.write(f"\nTable {i}:\n")
        for r, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            f.write(f"R{r}: {' | '.join(cells)}\n")
